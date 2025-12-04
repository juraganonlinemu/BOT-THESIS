import streamlit as st
import google.generativeai as genai
import pandas as pd
import requests
import re
import os
import time
import datetime
import PyPDF2
from docx import Document
from io import BytesIO

# ==========================================
# 1. KONFIGURASI & STYLE
# ==========================================
st.set_page_config(page_title="Thesis Master V4.4 (Multi-Discipline)", page_icon="🎓", layout="wide")

st.markdown("""
<style>
    .main-header {font-size: 2.2rem; font-weight: 800; color: #1E88E5; margin-bottom: 0px;}
    .sub-header {font-size: 1.1rem; color: #424242; font-style: italic;}
    .success-box {background-color: #E8F5E9; padding: 15px; border-radius: 8px; border-left: 5px solid #2E7D32;}
    .formula-box {background-color: #F3E5F5; padding: 10px; border-radius: 5px; border-left: 5px solid #9C27B0; margin-bottom: 10px;}
    .stTextArea textarea {font-size: 14px; line-height: 1.6; font-family: 'Times New Roman', serif;}
    .stButton>button {width: 100%; border-radius: 6px; font-weight: 600;}
</style>
""", unsafe_allow_html=True)

PDF_CACHE_FILE = "cached_pdf_context.txt"

# ==========================================
# 2. FILE OPERATIONS
# ==========================================
def save_to_disk(filename, content):
    try:
        with open(filename, "w", encoding="utf-8") as f: f.write(str(content))
    except: pass

def load_from_disk(filename):
    if os.path.exists(filename):
        with open(filename, "r", encoding="utf-8") as f: return f.read()
    return ""

def delete_file(filename):
    if os.path.exists(filename): os.remove(filename)

def perform_full_reset():
    delete_file(PDF_CACHE_FILE)
    keys_to_reset = ['naskah_bab1', 'naskah_bab2', 'naskah_bab3', 'naskah_bab4', 'naskah_bab5',
                     'outline_bab1', 'outline_bab2', 'outline_bab3', 'outline_bab4', 'outline_bab5', 
                     'judul_final', 'data_penelitian', 'search_results', 'generated_formulas', 'bidang_ilmu']
    for k in keys_to_reset:
        if k in st.session_state: st.session_state[k] = "" if 'outline' not in k and 'search' not in k else []
        delete_file(f"backup_{k}.txt")
    st.session_state['pdf_context'] = ""
    return True

# ==========================================
# 3. SEARCH ENGINE (FILTER TAHUN + EXCEL)
# ==========================================
def search_jurnal(keyword, limit=10, max_age_years=10):
    url = "https://api.crossref.org/works"
    current_year = datetime.datetime.now().year
    start_year = current_year - max_age_years
    filter_param = f"type:journal-article,from-pub-date:{start_year}-01-01"
    
    params = {
        "query": keyword,
        "rows": limit,
        "select": "title,DOI,URL,author,published-print,created",
        "filter": filter_param,
        "sort": "published",
        "order": "desc"
    }
    
    try:
        response = requests.get(url, params=params, timeout=15)
        data = response.json()['message']['items']
        results = []
        for item in data:
            title = item.get('title', ['No Title'])[0]
            doi = item.get('DOI', '-')
            link = item.get('URL', '-')
            try: year = item['published-print']['date-parts'][0][0]
            except: 
                try: year = item['created']['date-parts'][0][0]
                except: year = "-"
            try: auth = item['author'][0]['family'] + " et al."
            except: auth = "Unknown"

            results.append({"Judul": title, "Penulis": auth, "Tahun": year, "DOI": doi, "Link": link})
        return results
    except Exception as e: return []

def generate_search_formulas(title, bidang, api_key):
    if not title: return []
    p = f"""
    Bertindaklah sebagai Pustakawan Riset di bidang {bidang}. 
    Buatkan 3 Variasi String Pencarian untuk judul: "{title}"
    Output WAJIB format Python List of Strings.
    """
    try:
        res = get_model(api_key).generate_content(p).text
        return clean_parse_list(res)
    except: return []

# ==========================================
# 4. PDF PROCESSING
# ==========================================
def extract_text_from_pdfs(uploaded_files):
    combined_text = ""
    for pdf_file in uploaded_files:
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page_num in range(min(len(pdf_reader.pages), 50)): 
                extracted = pdf_reader.pages[page_num].extract_text()
                if extracted: text += extracted + "\n"
            combined_text += f"\n--- SUMBER: {pdf_file.name} ---\n{text}\n"
        except Exception as e: st.error(f"Gagal baca {pdf_file.name}: {e}")
    return combined_text

def retrieve_relevant_context(query, context_text, top_k=3):
    if not context_text or len(context_text) < 100: return ""
    chunk_size = 4000
    chunks = [context_text[i:i+chunk_size] for i in range(0, len(context_text), chunk_size)]
    keywords = [k.lower() for k in query.split() if len(k) > 4]
    if not keywords: return chunks[0] if chunks else ""
    
    scored_chunks = []
    for chunk in chunks:
        score = sum(chunk.lower().count(k) for k in keywords)
        scored_chunks.append((score, chunk))
    scored_chunks.sort(key=lambda x: x[0], reverse=True)
    return "\n...[POTONGAN RELEVAN]...\n".join([c[1] for c in scored_chunks[:top_k]])

# ==========================================
# 5. AI ENGINE
# ==========================================
def get_model(api_key):
    genai.configure(api_key=api_key)
    safe = [
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"}
    ]
    
    found_model_name = "gemini-pro"
    try:
        models = list(genai.list_models())
        for m in models:
            if 'generateContent' in m.supported_generation_methods:
                if 'flash' in m.name: found_model_name = m.name; break
                if 'pro' in m.name and found_model_name == "gemini-pro": found_model_name = m.name
    except: pass

    class AutoModel:
        def __init__(self, target): self.model_id = target
        def generate_content(self, prompt):
            try:
                model = genai.GenerativeModel(self.model_id, safety_settings=safe)
                return model.generate_content(prompt)
            except:
                return genai.GenerativeModel("gemini-pro", safety_settings=safe).generate_content(prompt)

    return AutoModel(found_model_name)

def clean_parse_list(text_response):
    try:
        match = re.search(r'\[.*\]', text_response, re.DOTALL)
        if match: return eval(match.group(0))
        return [l.strip().replace('- ','').replace('* ','') for l in text_response.split('\n') if l.strip()]
    except: return []

# ==========================================
# 6. WORD EXPORT
# ==========================================
def add_formatted_paragraph(doc, text, style='Normal'):
    p = doc.add_paragraph(style=style)
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**'): p.add_run(token[2:-2]).bold = True
        elif token.startswith('*') and token.endswith('*'): p.add_run(token[1:-1]).italic = True
        else: p.add_run(token)

def create_docx_from_text(text_content):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = 12 * 12700 
    for line in text_content.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('# '): doc.add_heading(line.replace('# ','').replace('**',''), 0)
        elif line.startswith('## '): doc.add_heading(line.replace('## ','').replace('**',''), 1)
        elif line.startswith('### '): doc.add_heading(line.replace('### ','').replace('**',''), 2)
        elif line.startswith('- ') or line.startswith('* '):
            clean_line = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', clean_line)
            for token in tokens:
                if token.startswith('**') and token.endswith('**'): p.add_run(token[2:-2]).bold = True
                elif token.startswith('*') and token.endswith('*'): p.add_run(token[1:-1]).italic = True
                else: p.add_run(token)
        else: add_formatted_paragraph(doc, line)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def convert_all_to_docx(data_naskah):
    full_text = ""
    for k in ['naskah_bab1', 'naskah_bab2', 'naskah_bab3', 'naskah_bab4', 'naskah_bab5']:
        if data_naskah.get(k): full_text += data_naskah[k] + "\n\n"
    return create_docx_from_text(full_text)

# ==========================================
# 7. STATE MANAGEMENT
# ==========================================
keys = ['naskah_bab1', 'naskah_bab2', 'naskah_bab3', 'naskah_bab4', 'naskah_bab5',
        'outline_bab1', 'outline_bab2', 'outline_bab3', 'outline_bab4', 'outline_bab5',
        'judul_final', 'pdf_context', 'data_penelitian', 'search_results', 'generated_formulas', 'bidang_ilmu']

for k in keys:
    if k not in st.session_state:
        if k == 'pdf_context': st.session_state[k] = load_from_disk(PDF_CACHE_FILE)
        elif 'outline' in k or 'search' in k or 'formula' in k: st.session_state[k] = []
        elif k == 'data_penelitian': st.session_state[k] = ""
        elif k == 'bidang_ilmu': st.session_state[k] = "Kesehatan/Keperawatan" # Default
        else: st.session_state[k] = load_from_disk(f"backup_{k}.txt")

# ==========================================
# 8. SIDEBAR
# ==========================================
with st.sidebar:
    st.markdown("## ⚙️ Control Panel")
    
    default_key = load_from_disk("my_key.txt")
    api_key = st.text_input("Gemini API Key", value=default_key, type="password")
    if st.checkbox("Simpan Key", value=True if default_key else False):
        if api_key: save_to_disk("my_key.txt", api_key)
    
    st.divider()
    if st.button("⚠️ RESET PROYEK (Mulai Baru)", type="primary"):
        perform_full_reset()
        st.rerun()
        
    st.divider()
    st.markdown("### 📚 1. Identitas & Bidang")
    # --- FITUR BIDANG ILMU DIKEMBALIKAN DISINI ---
    opsi_bidang = ["Kesehatan/Keperawatan", "Kedokteran", "Teknik (Informatika/Sipil/Mesin)", 
                   "Ekonomi & Bisnis", "Hukum", "Pendidikan/Keguruan", "Psikologi", "Sosial & Politik", "Pertanian"]
    
    # Load selection from state if exists
    idx = 0
    if st.session_state['bidang_ilmu'] in opsi_bidang:
        idx = opsi_bidang.index(st.session_state['bidang_ilmu'])
        
    pilih_bidang = st.selectbox("Pilih Bidang Keilmuan:", opsi_bidang, index=idx)
    if pilih_bidang != st.session_state['bidang_ilmu']:
        st.session_state['bidang_ilmu'] = pilih_bidang
        save_to_disk("backup_bidang_ilmu.txt", pilih_bidang)

    st.divider()
    st.markdown("### 📚 2. Upload Referensi")
    uploaded_pdfs = st.file_uploader("Upload Jurnal", type=['pdf'], accept_multiple_files=True)
    if uploaded_pdfs:
        if st.button("🔄 Proses PDF"):
            with st.spinner("Mengekstrak..."):
                raw = extract_text_from_pdfs(uploaded_pdfs)
                st.session_state['pdf_context'] += raw
                save_to_disk(PDF_CACHE_FILE, st.session_state['pdf_context'])
                st.success("Tersimpan!"); time.sleep(1); st.rerun()
    
    if len(st.session_state['pdf_context']) > 100:
        st.success(f"✅ PDF Aktif ({len(st.session_state['pdf_context'])} chars)")
    else:
        st.warning("⚠️ Belum ada PDF")

    st.divider()
    judul_input = st.text_area("Judul Tesis:", value=st.session_state['judul_final'], height=80)
    if judul_input != st.session_state['judul_final']:
        st.session_state['judul_final'] = judul_input
        save_to_disk("backup_judul_final.txt", judul_input)

# ==========================================
# 9. MAIN APP
# ==========================================
st.markdown('<div class="main-header">🎓 Thesis Master V4.4</div>', unsafe_allow_html=True)
st.markdown(f'<div class="sub-header">Mode: {st.session_state["bidang_ilmu"]} | Fitur Lengkap</div>', unsafe_allow_html=True)
st.divider()

if not api_key: st.stop()

tab_search, tab1, tab2, tab3, tab4, tab5, tab_dl = st.tabs(["🔎 Cari Jurnal", "BAB 1", "BAB 2", "BAB 3", "BAB 4", "BAB 5", "💾 Export All"])

# --- TAB PENCARIAN ---
with tab_search:
    st.header(f"🔎 Riset Jurnal ({st.session_state['bidang_ilmu']})")
    
    with st.expander("💡 Rekomendasi String Pencarian", expanded=True):
        if not st.session_state['judul_final']:
            st.info("⚠️ Masukkan Judul Tesis di Sidebar dulu.")
        else:
            if st.button("✨ Generate Rumus"):
                with st.spinner("Menganalisa..."):
                    formulas = generate_search_formulas(st.session_state['judul_final'], st.session_state['bidang_ilmu'], api_key)
                    st.session_state['generated_formulas'] = formulas
            
            if st.session_state['generated_formulas']:
                cols = st.columns(len(st.session_state['generated_formulas']))
                for i, f in enumerate(st.session_state['generated_formulas']):
                    with cols[i]:
                        st.markdown(f"<div class='formula-box'><b>Opsi {i+1}:</b><br>{f}</div>", unsafe_allow_html=True)
                        st.code(f, language="text")

    st.divider()
    col_s1, col_s2, col_s3 = st.columns([3, 1, 1])
    keyword = col_s1.text_input("Kata Kunci:", placeholder="Topik penelitian...")
    limit = col_s2.number_input("Jml Hasil", 5, 50, 10)
    max_age = col_s3.number_input("Max Umur (Thn)", 1, 20, 10)
    
    if st.button("Cari Jurnal"):
        with st.spinner(f"Mencari referensi {max_age} tahun terakhir..."):
            res = search_jurnal(keyword, limit, max_age_years=max_age)
            st.session_state['search_results'] = res
            if not res: st.warning("Nihil.")
            else: st.success(f"Ditemukan {len(res)} jurnal!")

    if st.session_state['search_results']:
        df = pd.DataFrame(st.session_state['search_results'])
        st.dataframe(df, hide_index=True)
        c_csv, c_xls = st.columns(2)
        c_csv.download_button("📥 Download CSV", df.to_csv(index=False).encode('utf-8'), "jurnal.csv", "text/csv")
        try:
            output = BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer: df.to_excel(writer, index=False)
            c_xls.download_button("📥 Download Excel", output.getvalue(), "jurnal.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        except: c_xls.error("Install openpyxl utk Excel")

# --- RENDER FUNCTION ---
def render_chapter(key_naskah, key_outline, label, extra_context=""):
    col1, col2 = st.columns([1, 3])
    widget_key = f"area_{key_naskah}"
    if widget_key in st.session_state:
        if st.session_state[widget_key] != st.session_state[key_naskah]:
            st.session_state[key_naskah] = st.session_state[widget_key]
            save_to_disk(f"backup_{key_naskah}.txt", st.session_state[key_naskah])

    with col1:
        st.markdown(f"#### Struktur {label}")
        if st.button(f"Buat Outline", key=f"btn_{key_naskah}"):
            if not st.session_state['judul_final']: st.error("Judul kosong!"); return
            with st.spinner("Membuat outline..."):
                p = f"""
                Buatkan Outline Detail {label} Tesis: '{st.session_state['judul_final']}'. 
                Bidang Keilmuan: {st.session_state['bidang_ilmu']}.
                Output list Python string.
                """
                try:
                    res = get_model(api_key).generate_content(p).text
                    st.session_state[key_outline] = clean_parse_list(res)
                except Exception as e: st.error(str(e))
        
        if st.session_state[key_outline]:
            txt = st.text_area("Edit Outline:", "\n".join(st.session_state[key_outline]), height=250, key=f"txt_{key_outline}")
            st.session_state[key_outline] = [x.strip() for x in txt.split('\n') if x.strip()]
        
        st.divider()
        if st.session_state[key_naskah]:
            docx_bab = create_docx_from_text(st.session_state[key_naskah])
            st.download_button(f"📥 Download {label}", docx_bab, f"{label}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    with col2:
        st.markdown(f"#### ✍️ Editor {label}")
        selected_sub = ""
        if st.session_state[key_outline]:
            selected_sub = st.selectbox("Pilih Sub-Bab:", st.session_state[key_outline], key=f"sel_{key_naskah}")
            if st.button(f"🚀 Tulis: {selected_sub}", key=f"go_{key_naskah}", type="primary"):
                with st.spinner(f"Menulis {selected_sub}..."):
                    ctx = retrieve_relevant_context(selected_sub, st.session_state['pdf_context'])
                    p = f"""
                    Anda adalah Ahli Riset Akademik di bidang: {st.session_state['bidang_ilmu']}.
                    Tulis Tesis bagian: "{selected_sub}". 
                    Judul: {st.session_state['judul_final']}
                    Konteks PDF: {ctx}. Info: {extra_context}
                    ATURAN: FOKUS sub-bab ini. Panjang 500-800 kata. Gaya Bahasa Akademik bidang {st.session_state['bidang_ilmu']}.
                    Sitasi (Author, Tahun). Gunakan **bold** untuk penekanan, *italic* istilah asing.
                    """
                    try:
                        res = get_model(api_key).generate_content(p).text
                        if res:
                            header = f"\n\n## {selected_sub}\n"
                            st.session_state[key_naskah] += header + res
                            st.session_state[widget_key] = st.session_state[key_naskah]
                            save_to_disk(f"backup_{key_naskah}.txt", st.session_state[key_naskah])
                            st.success("Selesai!"); time.sleep(0.5); st.rerun()
                    except Exception as e: st.error(f"Error: {e}")

        st.text_area("Isi Naskah:", value=st.session_state[key_naskah], height=500, key=widget_key)

with tab1: render_chapter('naskah_bab1', 'outline_bab1', 'BAB 1')
with tab2: render_chapter('naskah_bab2', 'outline_bab2', 'BAB 2')
with tab3: render_chapter('naskah_bab3', 'outline_bab3', 'BAB 3')
with tab4:
    st.header("📊 BAB 4: HASIL & PEMBAHASAN")
    col_d1, col_d2 = st.columns(2)
    with col_d1:
        if st.button("🎲 Generate Dummy Data"):
            dummy = "Tabel 1. Distribusi Data\n- Variabel X: 40%\n- Variabel Y: 60%\n- Signifikansi: 0.003"
            st.session_state['data_penelitian'] = dummy
            st.rerun()
    data_input = st.text_area("Input Data:", value=st.session_state['data_penelitian'], height=150)
    st.session_state['data_penelitian'] = data_input
    st.divider()
    render_chapter('naskah_bab4', 'outline_bab4', 'BAB 4', extra_context=f"DATA: {st.session_state['data_penelitian']}")
with tab5:
    st.header("🏁 BAB 5: KESIMPULAN")
    ringkasan_bab4 = st.session_state['naskah_bab4'][:2000] if st.session_state['naskah_bab4'] else ""
    render_chapter('naskah_bab5', 'outline_bab5', 'BAB 5', extra_context=f"Berdasarkan: {ringkasan_bab4}")

with tab_dl:
    st.header("💾 Download Lengkap")
    if st.button("📥 DOWNLOAD FULL TESIS"):
        docx = convert_all_to_docx(st.session_state)
        st.download_button("Klik Unduh Word", docx, "Tesis_Full.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")