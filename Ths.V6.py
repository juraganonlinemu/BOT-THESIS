import streamlit as st
import google.generativeai as genai
import pandas as pd
import requests
import xml.etree.ElementTree as ET
import re
import os
import time
import datetime
import PyPDF2
from docx import Document
from io import BytesIO

# ==========================================
# 1. KONFIGURASI SISTEM
# ==========================================
st.set_page_config(page_title="Thesis Master V6.2 (Stabil)", page_icon="🎓", layout="wide")

# --- KOLAM API KEY ---
API_KEYS_POOL = [
    "KOSONGKAN_SAJA_INI_UNTUK_GITHUB", 
    "KARENA_KEY_ASLI_ADA_DI_SECRETS_STREAMLIT"
]

try:
    if "api_keys" in st.secrets:
        API_KEYS_POOL = st.secrets["api_keys"]
except: pass

st.markdown("""
<style>
    .main-header {font-size: 2.2rem; font-weight: 800; color: #1E88E5; margin-bottom: 0px;}
    .sub-header {font-size: 1.1rem; color: #424242; font-style: italic;}
    .stTextArea textarea {font-size: 14px; line-height: 1.6; font-family: 'Times New Roman', serif;}
    .stButton>button {width: 100%; border-radius: 6px; font-weight: 600;}
    .formula-box {
        background-color: #e3f2fd;
        border-left: 5px solid #1565c0;
        padding: 15px;
        margin-bottom: 10px;
        border-radius: 5px;
        font-family: monospace;
        font-size: 0.85rem;
        color: #000;
    }
</style>
""", unsafe_allow_html=True)

PDF_CACHE_FILE = "cached_pdf_context.txt"

# ==========================================
# 2. SISTEM LOGIN
# ==========================================
def check_access():
    try: DATABASE_USER = st.secrets["pengguna"]
    except:
        # User Default (Hanya contoh dummy, bukan data mahasiswa asli)
        DATABASE_USER = {
            "CONTOH-USER": "2025-01-01"
        }

    if st.session_state.get('is_logged_in', False): return True

    st.markdown("## 🔒 Restricted Access")
    col1, col2 = st.columns([2, 1])
    with col1: token_input = st.text_input("Kode Akses:", type="password")
    
    if st.button("Masuk"):
        if token_input in DATABASE_USER:
            expiry_str = DATABASE_USER[token_input]
            try: expiry = datetime.datetime.strptime(expiry_str, "%Y-%m-%d").date()
            except: expiry = datetime.date.today() + datetime.timedelta(days=365)

            if datetime.date.today() <= expiry:
                st.session_state['is_logged_in'] = True
                st.session_state['username'] = token_input
                st.success("Login Berhasil!"); time.sleep(0.5); st.rerun()
            else: st.error("Kode Kadaluarsa.")
        else: st.error("Kode Salah.")
    return False

if not check_access(): st.stop()

# ==========================================
# 3. FILE OPERATIONS
# ==========================================
def get_user_filename(filename):
    user = st.session_state.get('username', 'guest')
    safe_user = re.sub(r'[^a-zA-Z0-9]', '_', user)
    return f"{safe_user}_{filename}"

def save_to_disk(filename, content):
    try:
        with open(get_user_filename(filename), "w", encoding="utf-8") as f: f.write(str(content))
    except: pass

def load_from_disk(filename):
    unique = get_user_filename(filename)
    if os.path.exists(unique):
        with open(unique, "r", encoding="utf-8") as f: return f.read()
    return ""

def delete_file(filename):
    unique = get_user_filename(filename)
    if os.path.exists(unique): os.remove(unique)

def perform_full_reset():
    delete_file(PDF_CACHE_FILE)
    keys = ['naskah_bab1', 'naskah_bab2', 'naskah_bab3', 'naskah_bab4', 'naskah_bab5',
            'outline_bab1', 'outline_bab2', 'outline_bab3', 'outline_bab4', 'outline_bab5', 
            'judul_final', 'data_penelitian', 'search_results', 'opsi_judul_generated', 'generated_formulas']
    for k in keys:
        if k in st.session_state: st.session_state[k] = "" if 'outline' not in k and 'search' not in k else []
        delete_file(f"backup_{k}.txt")
    st.session_state['pdf_context'] = ""
    return True

# ==========================================
# 4. AI ENGINE
# ==========================================
def get_model():
    safe = [
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"}
    ]
    class RotatingModel:
        def generate_content(self, prompt):
            errors = []
            for key in API_KEYS_POOL:
                if "MASUKKAN" in key: continue 
                try:
                    genai.configure(api_key=key)
                    found = "gemini-pro"
                    try:
                        for m in genai.list_models():
                            if 'generateContent' in m.supported_generation_methods:
                                if 'flash' in m.name: found = m.name; break
                    except: pass
                    return genai.GenerativeModel(found, safety_settings=safe).generate_content(prompt)
                except Exception as e:
                    if "429" in str(e) or "quota" in str(e).lower(): continue 
                    else: errors.append(str(e))
            raise Exception(f"Semua API Key sibuk/limit. Detail: {errors}")
    return RotatingModel()

def clean_parse_list(text_response):
    try:
        match = re.search(r'\[.*\]', text_response, re.DOTALL)
        if match: 
            parsed = eval(match.group(0))
            flat = []
            def flatten(items):
                for x in items:
                    if isinstance(x, list): flatten(x)
                    else: flat.append(str(x).strip())
            flatten(parsed)
            return flat
        return [l.strip().replace('- ','').replace('* ','') for l in text_response.split('\n') if l.strip()]
    except: return []

# ==========================================
# 5. SMART SEARCH
# ==========================================
def search_pubmed(keyword, limit=10, max_age_years=10, is_fulltext=False):
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"
    year_now = datetime.date.today().year
    min_year = year_now - max_age_years
    
    term = f"{keyword} AND {min_year}:{year_now}[pdat]"
    if is_fulltext: term += " AND free full text[sb]"
        
    try:
        search_url = f"{base_url}/esearch.fcgi?db=pubmed&term={term}&retmax={limit*2}&retmode=json"
        res = requests.get(search_url, timeout=10).json()
        id_list = res.get('esearchresult', {}).get('idlist', [])
        
        if not id_list: return []
        
        ids = ','.join(id_list)
        fetch_url = f"{base_url}/efetch.fcgi?db=pubmed&id={ids}&retmode=xml"
        root = ET.fromstring(requests.get(fetch_url, timeout=15).content)
        
        results = []
        for article in root.findall(".//PubmedArticle"):
            try:
                title = article.find(".//ArticleTitle").text
                try: year = int(article.find(".//PubDate/Year").text)
                except: year = year_now
                
                if year > year_now or year < min_year: continue

                try: auth = article.find(".//AuthorList/Author/LastName").text + " et al."
                except: auth = "Team"
                
                pmid = article.find(".//PMID").text
                doi = "-"
                for i in article.findall(".//ArticleIdList/ArticleId"):
                    if i.get('IdType') == 'doi': doi = i.text
                
                link = f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/"
                results.append({"Source": "PubMed", "Judul": title, "Penulis": auth, "Tahun": year, "DOI": doi, "Link": link})
                if len(results) >= limit: break
            except: continue
        return results
    except: return []

def search_crossref(keyword, limit=10, max_age_years=10, is_oa=False, is_fulltext=False):
    url = "https://api.crossref.org/works"
    year_now = datetime.date.today().year
    min_year = year_now - max_age_years
    
    clean_kw = re.sub(r'[^\x00-\x7F]+', '', keyword)
    
    params = {
        "query": clean_kw, 
        "rows": limit * 5,
        "select": "title,DOI,URL,author,published-print,published-online,created,link,license",
        "filter": f"type:journal-article,from-pub-date:{min_year}-01-01",
        "sort": "published", "order": "desc"
    }
    
    try:
        data = requests.get(url, params=params, timeout=20).json()['message']['items']
        final_results = []
        for i in data:
            y = 0
            if 'published-print' in i: y = i['published-print']['date-parts'][0][0]
            elif 'published-online' in i: y = i['published-online']['date-parts'][0][0]
            elif 'created' in i: y = i['created']['date-parts'][0][0]
            if y > year_now or y < min_year: continue

            l = i.get('URL', '-')
            pdf_link = ""
            has_pdf = False
            if 'link' in i:
                for lnk in i['link']:
                    if 'pdf' in lnk.get('content-type', '').lower():
                        pdf_link = lnk.get('URL'); has_pdf = True; break
            if not pdf_link: pdf_link = l 

            has_license = 'license' in i
            if is_fulltext and not has_pdf: continue
            if is_oa and not has_license: continue

            t = i.get('title', ['No Title'])[0]
            d = i.get('DOI', '-')
            try: a = i['author'][0]['family'] + " et al."
            except: a = "Unknown"

            final_results.append({"Source": "Crossref", "Judul": t, "Penulis": a, "Tahun": y, "DOI": d, "Link": pdf_link})
            if len(final_results) >= limit: break
        return final_results
    except: return []

def smart_search_dispatcher(bidang, kw, lim, age, oa, ft):
    results = []
    if "Kesehatan" in bidang or "Kedokteran" in bidang:
        results += search_pubmed(kw, lim, age, ft)
    results += search_crossref(kw, lim, age, oa, ft)
    seen = set(); unique = []
    for r in results:
        if r['Judul'] not in seen:
            unique.append(r); seen.add(r['Judul'])
    return unique[:lim]

def generate_search_formulas(title, bidang):
    if not title: return []
    p = f"""
    Bertindaklah sebagai Pustakawan Riset Senior.
    Judul Penelitian: "{title}" (Bidang: {bidang}).
    
    TUGAS: Buatkan 3 Rumus Pencarian (Search String) untuk database jurnal internasional (PubMed/Scopus).
    
    Output WAJIB berupa List Python ['Rumus 1', 'Rumus 2', 'Rumus 3'].
    
    Kriteria Rumus:
    1. Rumus Basic: Keyword utama saja.
    2. Rumus Menengah: Gunakan sinonim sederhana.
    3. Rumus ADVANCED BOOLEAN: Gunakan tanda kurung, OR untuk sinonim, dan AND antar variabel.
       Contoh Format: ("Var1" OR "Sinonim1") AND ("Var2" OR "Sinonim2") AND ("Var3")
    """
    try: return clean_parse_list(get_model().generate_content(p).text)
    except: return []

def extract_text_from_pdfs(files):
    text = ""
    for f in files:
        try:
            reader = PyPDF2.PdfReader(f)
            for p in range(min(len(reader.pages), 50)):
                text += reader.pages[p].extract_text() + "\n"
            text += f"\n--- SUMBER: {f.name} ---\n"
        except: pass
    return text

def retrieve_relevant_context(query, text, top_k=3):
    if not text: return ""
    chunks = [text[i:i+4000] for i in range(0, len(text), 4000)]
    scored = []
    keywords = [k.lower() for k in query.split() if len(k)>4]
    for c in chunks:
        score = sum(c.lower().count(k) for k in keywords)
        scored.append((score, c))
    scored.sort(key=lambda x: x[0], reverse=True)
    return "\n...".join([s[1] for s in scored[:top_k]])

# ==========================================
# 6. EXPORT WORD
# ==========================================
def create_docx(text):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'; style.font.size = 12 * 12700
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('# '): doc.add_heading(line.replace('# ','').replace('**',''), 0)
        elif line.startswith('## '): doc.add_heading(line.replace('## ','').replace('**',''), 1)
        elif line.startswith('### '): doc.add_heading(line.replace('### ','').replace('**',''), 2)
        elif line.startswith('- '): 
            p = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line[2:])
            for pt in parts:
                if pt.startswith('**'): p.add_run(pt[2:-2]).bold = True
                elif pt.startswith('*'): p.add_run(pt[1:-1]).italic = True
                else: p.add_run(pt)
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
            for pt in parts:
                if pt.startswith('**'): p.add_run(pt[2:-2]).bold = True
                elif pt.startswith('*'): p.add_run(pt[1:-1]).italic = True
                else: p.add_run(pt)
    bio = BytesIO(); doc.save(bio); return bio.getvalue()

def convert_all_to_docx(state):
    full = ""
    for k in ['naskah_bab1', 'naskah_bab2', 'naskah_bab3', 'naskah_bab4', 'naskah_bab5']:
        if state.get(k): full += state[k] + "\n\n"
    return create_docx(full)

# ==========================================
# 7. LOAD STATE (BUG FIXED HERE)
# ==========================================
keys = ['naskah_bab1', 'naskah_bab2', 'naskah_bab3', 'naskah_bab4', 'naskah_bab5',
        'outline_bab1', 'outline_bab2', 'outline_bab3', 'outline_bab4', 'outline_bab5',
        'judul_final', 'pdf_context', 'data_penelitian', 'search_results', 
        'opsi_judul_generated', 'generated_formulas', 'bidang_ilmu'] # <-- 'generated_formulas' ADDED

for k in keys:
    if k not in st.session_state:
        if k == 'pdf_context': st.session_state[k] = load_from_disk(PDF_CACHE_FILE)
        elif 'outline' in k or 'search' in k or 'opsi' in k or 'generated' in k: st.session_state[k] = []
        elif k == 'bidang_ilmu': st.session_state[k] = "Kesehatan/Keperawatan"
        else: st.session_state[k] = load_from_disk(f"backup_{k}.txt")

# ==========================================
# 8. UI & MAIN
# ==========================================
with st.sidebar:
    st.markdown("## ⚙️ Control Panel")
    if st.button("⚠️ RESET PROYEK", type="primary"): perform_full_reset(); st.rerun()
    st.divider()
    
    opsi_bidang = ["Kesehatan/Keperawatan", "Kedokteran", "Teknik", "Ekonomi", "Hukum", "Pendidikan", "Sosial", "Pertanian"]
    idx = opsi_bidang.index(st.session_state['bidang_ilmu']) if st.session_state['bidang_ilmu'] in opsi_bidang else 0
    bidang = st.selectbox("Bidang:", opsi_bidang, index=idx)
    if bidang != st.session_state['bidang_ilmu']: 
        st.session_state['bidang_ilmu'] = bidang
        save_to_disk("backup_bidang_ilmu.txt", bidang)

    st.markdown("---")
    ide = st.text_area("Topik Awal:", height=70)
    if st.button("✨ Rekomendasi Judul"):
        if ide:
            with st.spinner("..."):
                p = f"Buat 3 Judul Tesis {bidang} topik {ide}. Output List Python. TANPA BASA BASI."
                try: st.session_state['opsi_judul_generated'] = clean_parse_list(get_model().generate_content(p).text)
                except Exception as e: st.error(str(e))
    if st.session_state['opsi_judul_generated']:
        pilihan = st.radio("Pilih:", st.session_state['opsi_judul_generated'])
        if st.button("Pakai"): st.session_state['judul_final'] = pilihan; st.rerun()

    judul_input = st.text_area("Judul Final:", value=st.session_state['judul_final'], height=100)
    if judul_input != st.session_state['judul_final']:
        st.session_state['judul_final'] = judul_input
        save_to_disk("backup_judul_final.txt", judul_input)

    st.divider()
    pdfs = st.file_uploader("Upload PDF (Jurnal Acuan)", type=['pdf'], accept_multiple_files=True)
    if pdfs and st.button("Proses PDF"):
        with st.spinner("Membaca..."):
            raw = extract_text_from_pdfs(pdfs)
            st.session_state['pdf_context'] += raw
            save_to_disk(PDF_CACHE_FILE, st.session_state['pdf_context'])
            st.success("Sukses! Data PDF tersimpan.")
    if len(st.session_state['pdf_context']) > 100: st.success("✅ PDF Ready")

st.markdown(f'<div class="main-header">🎓 Thesis Master V6.2</div>', unsafe_allow_html=True)
st.markdown(f'<div class="sub-header">Mode: {st.session_state["bidang_ilmu"]} | Bug Fixed</div>', unsafe_allow_html=True)
st.divider()

tabs = st.tabs(["🔎 Riset", "BAB 1", "BAB 2", "BAB 3", "BAB 4", "BAB 5", "💾 Export"])

with tabs[0]:
    st.header("🔎 Pencarian Jurnal")
    
    with st.expander("💡 Rekomendasi Rumus Boolean (Klik Disini)"):
        if st.button("✨ Generate Rumus Canggih"):
            if st.session_state['judul_final']:
                with st.spinner("Meracik rumus..."):
                    forms = generate_search_formulas(st.session_state['judul_final'], st.session_state['bidang_ilmu'])
                    st.session_state['generated_formulas'] = forms
        
        if st.session_state.get('generated_formulas'):
            for f in st.session_state['generated_formulas']:
                st.markdown(f"<div class='formula-box'>{f}</div>", unsafe_allow_html=True)

    c1, c2, c3 = st.columns([2,1,1])
    kw = c1.text_input("Keyword / Formula:", placeholder="Paste rumus Boolean di sini...")
    lim = c2.number_input("Limit Hasil", 5, 50, 10)
    age = c3.number_input("Rentang Thn", 1, 20, 10)
    
    c_oa, c_ft = st.columns(2)
    oa_only = c_oa.checkbox("Hanya Open Access")
    ft_only = c_ft.checkbox("Hanya Fulltext Link")
    
    if st.button("Cari Jurnal"):
        with st.spinner(f"Mencari data..."):
            st.session_state['search_results'] = smart_search_dispatcher(st.session_state['bidang_ilmu'], kw, lim, age, oa_only, ft_only)
            
    if st.session_state['search_results']:
        df = pd.DataFrame(st.session_state['search_results'])
        if df.empty:
            st.warning("Tidak ditemukan. Coba keyword yang lebih sederhana.")
        else:
            st.success(f"Ditemukan {len(df)} jurnal.")
            st.dataframe(df, hide_index=True)
            try:
                out = BytesIO(); 
                with pd.ExcelWriter(out, engine='openpyxl') as w: df.to_excel(w, index=False)
                st.download_button("📥 Excel", out.getvalue(), "jurnal.xlsx")
            except: st.error("Install openpyxl")

def render_bab(key_naskah, key_outline, label, extra=""):
    c1, c2 = st.columns([1,3])
    wk = f"w_{key_naskah}"
    if wk in st.session_state and st.session_state[wk] != st.session_state[key_naskah]:
        st.session_state[key_naskah] = st.session_state[wk]
        save_to_disk(f"backup_{key_naskah}.txt", st.session_state[key_naskah])

    with c1:
        st.markdown(f"**Struktur {label}**")
        if st.button(f"Buat Outline", key=f"b_{key_naskah}"):
            if not st.session_state['judul_final']: st.error("Isi Judul dulu!"); return
            with st.spinner("..."):
                p = f"Buat Outline {label} Tesis: '{st.session_state['judul_final']}'. Bidang: {st.session_state['bidang_ilmu']}. List Python."
                try: st.session_state[key_outline] = clean_parse_list(get_model().generate_content(p).text)
                except Exception as e: st.error(str(e))
        if st.session_state[key_outline]:
            t = st.text_area("Edit Outline:", "\n".join(st.session_state[key_outline]), height=300, key=f"t_{key_outline}")
            st.session_state[key_outline] = [x.strip() for x in t.split('\n') if x.strip()]
        
        st.divider()
        if st.button(f"🗑️ Reset {label}", key=f"rst_{key_naskah}", type="secondary"):
            st.session_state[key_naskah] = ""
            save_to_disk(f"backup_{key_naskah}.txt", "")
            st.rerun()

    with c2:
        st.markdown(f"**Editor {label}**")
        if st.session_state[key_naskah]:
            st.download_button(f"📥 DOWNLOAD {label}", create_docx(st.session_state[key_naskah]), f"{label}.docx", type="primary")
        
        st.divider()
        sub = ""
        if st.session_state[key_outline]:
            c_sel, c_btn = st.columns([3,1])
            with c_sel: sub = st.selectbox("Pilih Sub-Bab:", st.session_state[key_outline], key=f"s_{key_naskah}")
            with c_btn: 
                if st.button(f"🚀 Tulis", key=f"g_{key_naskah}"):
                    with st.spinner("Menulis..."):
                        ctx = retrieve_relevant_context(sub, st.session_state['pdf_context'])
                        
                        p = f"""
                        Peran: Penulis Tesis Akademik ({st.session_state['bidang_ilmu']}).
                        Topik Sub-Bab: '{sub}'. Judul: {st.session_state['judul_final']}.
                        
                        DATA PENDUKUNG (CONTEXT):
                        {ctx}
                        {extra}
                        
                        INSTRUKSI KHUSUS (WAJIB DIPATUHI):
                        1. FILTER TAHUN CERDAS: Jika ada data di 'Context' yang tahunnya < 2015, JANGAN jadikan argumen utama.
                        2. SITASI CERDAS: Gunakan format APA (Nama, Tahun). DILARANG KERAS menggunakan angka [1].
                        3. GAYA BAHASA HUMANIS: Tulis dengan gaya natural manusia, variasi kalimat, hindari kata robotik.
                        4. NO PROLOG.
                        5. Panjang: 600-900 kata.
                        """
                        try:
                            res = get_model().generate_content(p).text
                            st.session_state[key_naskah] += f"\n\n## {sub}\n{res}"
                            st.session_state[wk] = st.session_state[key_naskah]
                            save_to_disk(f"backup_{key_naskah}.txt", st.session_state[key_naskah])
                            st.success("Selesai!"); time.sleep(0.5); st.rerun()
                        except Exception as e: st.error(str(e))
        
        st.text_area("Isi Naskah:", value=st.session_state[key_naskah], height=600, key=wk)

with tabs[1]: render_bab('naskah_bab1', 'outline_bab1', 'BAB 1')
with tabs[2]: render_bab('naskah_bab2', 'outline_bab2', 'BAB 2')
with tabs[3]: render_bab('naskah_bab3', 'outline_bab3', 'BAB 3')
with tabs[4]: 
    col_d1, col_d2 = st.columns(2)
    with col_d1:
        if st.button("🎲 Data Dummy"): st.session_state['data_penelitian'] = "Tabel 1: X=40%, Y=60%, p=0.003"; st.rerun()
    d_in = st.text_area("Data:", st.session_state['data_penelitian'])
    st.session_state['data_penelitian'] = d_in
    render_bab('naskah_bab4', 'outline_bab4', 'BAB 4', f"DATA: {d_in}")
with tabs[5]: render_bab('naskah_bab5', 'outline_bab5', 'BAB 5', f"Ringkasan: {st.session_state['naskah_bab4'][:1000]}")
with tabs[6]:
    st.download_button("📥 DOWNLOAD FULL TESIS", convert_all_to_docx(st.session_state), "Tesis_Full.docx", type="primary")